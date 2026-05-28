"""
文档处理与向量化模块 (RAG)
负责：加载文档 → 分块 → 向量化 → 存入 ChromaDB → 检索

优化:
- [#9] RAG 检索质量提升：混合检索（向量 + BM25关键词） + 重排序
- [#10] 引用溯源：返回结果标注文档名 + 段落位置 + chunk_id
- [#11] Embedding 降级：当 Embedding API 不可用时（403/余额不足/网络错误），
       自动切换为关键词索引模式，文档直接保存到磁盘并支持关键词搜索
"""
import os
import re
import json
import logging
from typing import Optional

from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma

from app.config import settings

logger = logging.getLogger(__name__)

# 本地 Embedding 批量大小（本地模型无API限制，可适当增大）
EMBEDDING_BATCH_SIZE = 100

# ===== 单例模式：复用 Embedding 和 ChromaDB 连接 =====
_embeddings_instance = None
_vector_store_cache = {}  # agent_id -> ChromaDB instance（按智能体隔离）

# ===== [#11] Embedding 可用性标志 =====
# None = 尚未检测，True = 可用，False = 不可用（自动降级为关键词模式）
_embedding_available = None

# 全局知识库的 collection 名称
GLOBAL_COLLECTION_NAME = "langchain"

# Embedding 提供者配置："local" 使用本地免费模型，"openai" 使用云端API（消耗额度）
EMBEDDING_PROVIDER = os.environ.get("EMBEDDING_PROVIDER", "local")
LOCAL_EMBEDDING_MODEL = os.environ.get("LOCAL_EMBEDDING_MODEL", "BAAI/bge-large-zh-v1.5")

# ===== [#11] 关键词索引配置 =====
KEYWORD_INDEX_DIR = os.path.join(os.path.dirname(settings.CHROMA_DIR) if hasattr(settings, 'CHROMA_DIR') else os.path.join(settings.DATA_DIR, 'keyword_index'), 'keyword_index') if hasattr(settings, 'DATA_DIR') else os.path.join(os.path.dirname(settings.CHROMA_DIR), 'keyword_index')

# 中文停用词
_STOPWORDS = {'的', '了', '是', '在', '和', '与', '有', '什么', '怎么', '如何', '哪些', '这个', '那个',
              '一个', '不是', '没有', '可以', '就是', '已经', '我们', '他们', '她们', '它们',
              '但是', '而且', '或者', '因为', '所以', '如果', '虽然', '而且', '以及',
              'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
              'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
              'should', 'may', 'might', 'shall', 'can', 'need', 'dare', 'ought',
              'used', 'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by', 'from'}


def get_indexing_mode() -> str:
    """获取当前索引模式（供外部查询，如健康检查）

    Returns:
        str: "vector"（向量模式）、"keyword"（关键词模式）、"unknown"（尚未检测）
    """
    global _embedding_available
    if _embedding_available is None:
        return "unknown"
    return "vector" if _embedding_available else "keyword"


def _get_keyword_index_path(agent_id: str = None) -> str:
    """获取关键词索引JSON文件路径

    Args:
        agent_id: 智能体ID，为None时使用全局索引

    Returns:
        str: JSON文件路径
    """
    os.makedirs(KEYWORD_INDEX_DIR, exist_ok=True)
    cache_key = agent_id or "__global__"
    safe_key = cache_key.replace('-', '_').replace(' ', '_')
    return os.path.join(KEYWORD_INDEX_DIR, f"index_{safe_key}.json")


def _load_keyword_index(agent_id: str = None) -> list[dict]:
    """从磁盘加载关键词索引

    Returns:
        list[dict]: 索引条目列表，每条包含 content, source_file, chunk_index
    """
    index_path = _get_keyword_index_path(agent_id)
    if not os.path.exists(index_path):
        return []
    try:
        with open(index_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.warning(f"加载关键词索引失败: {e}")
        return []


def _save_keyword_index(index_data: list[dict], agent_id: str = None):
    """保存关键词索引到磁盘

    Args:
        index_data: 索引条目列表
        agent_id: 智能体ID
    """
    index_path = _get_keyword_index_path(agent_id)
    os.makedirs(os.path.dirname(index_path), exist_ok=True)
    try:
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"保存关键词索引失败: {e}")


def _add_chunks_to_keyword_index(chunks: list, filename: str, agent_id: str = None):
    """将文档分块添加到关键词索引

    Args:
        chunks: LangChain Document 分块列表
        filename: 文档文件名
        agent_id: 智能体ID
    """
    index_data = _load_keyword_index(agent_id)

    # 先删除该文档的旧条目（避免重复）
    index_data = [entry for entry in index_data if entry.get("source_file") != filename]

    # 添加新条目
    for chunk in chunks:
        index_data.append({
            "content": chunk.page_content,
            "source_file": filename,
            "chunk_index": chunk.metadata.get("chunk_index", 0),
        })

    _save_keyword_index(index_data, agent_id)
    logger.info(f"关键词索引已更新: {filename}, 新增 {len(chunks)} 个分块, 索引总条目={len(index_data)}")


def _delete_from_keyword_index(filename: str, agent_id: str = None) -> int:
    """从关键词索引中删除指定文档的所有条目

    Args:
        filename: 文档文件名
        agent_id: 智能体ID

    Returns:
        int: 被删除的条目数
    """
    index_data = _load_keyword_index(agent_id)
    original_count = len(index_data)
    index_data = [entry for entry in index_data if entry.get("source_file") != filename]
    deleted_count = original_count - len(index_data)

    if deleted_count > 0:
        _save_keyword_index(index_data, agent_id)
        logger.info(f"关键词索引已删除: {filename}, 删除 {deleted_count} 个条目")

    return deleted_count


def _search_keyword_index(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """[#11] 纯关键词检索（在关键词索引上搜索）

    不依赖 ChromaDB/Embedding，直接在 JSON 索引上做关键词匹配。
    当 Embedding API 不可用时作为主要检索手段。

    Args:
        query: 用户查询
        top_k: 返回最相关的K个结果
        agent_id: 智能体ID

    Returns:
        list[dict]: 检索结果列表
    """
    index_data = _load_keyword_index(agent_id)
    if not index_data:
        return []

    # 提取查询关键词
    query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
    query_terms = query_terms - _STOPWORDS

    if not query_terms:
        # 如果所有词都是停用词，则用原文做简单包含匹配
        query_terms = {query.lower()}

    scored = []
    for entry in index_data:
        content = entry.get("content", "")
        content_lower = content.lower()

        # 计算关键词匹配度
        match_count = sum(1 for term in query_terms if term in content_lower)
        if match_count == 0:
            continue

        # TF 近似得分：匹配关键词数 / 内容长度 * 归一化因子
        # 同时考虑匹配比例（匹配了多少个查询词）
        term_coverage = match_count / max(len(query_terms), 1)
        tf_score = match_count / max(len(content), 1) * 1000
        combined_score = tf_score * 0.6 + term_coverage * 100 * 0.4

        scored.append({
            "content": content,
            "source": entry.get("source_file", "未知来源"),
            "chunk_index": entry.get("chunk_index", -1),
            "relevance_score": round(combined_score, 4),
        })

    # 按得分排序
    scored.sort(key=lambda x: x["relevance_score"], reverse=True)
    return scored[:top_k]


def _is_embedding_error(e: Exception) -> bool:
    """判断异常是否为 Embedding 不可用错误

    包括：403 权限不足、429 余额不足、网络连接错误等
    """
    error_str = str(e).lower()
    # 403 权限/认证错误
    if '403' in error_str or 'no access' in error_str or 'forbidden' in error_str:
        return True
    # 429 余额/限流错误
    if '429' in error_str or '余额' in error_str or 'rate limit' in error_str or 'quota' in error_str:
        return True
    # 1113 智谱余额不足
    if '1113' in error_str:
        return True
    # 连接错误
    if 'connection' in error_str or 'timeout' in error_str or 'connect' in error_str:
        return True
    # API key 错误
    if 'api_key' in error_str or 'api key' in error_str or 'unauthorized' in error_str or 'invalid api' in error_str:
        return True
    # new_api_error（智谱特有）
    if 'new_api_error' in error_str:
        return True
    return False


def get_embeddings():
    """获取 Embedding 模型（单例复用，避免重复初始化）

    支持两种模式：
    - local（默认）：使用本地 BAAI/bge-large-zh-v1.5，永久免费，离线可用，中文质量顶级
    - openai：使用智谱 embedding-3 云端API，消耗额度，需联网

    通过环境变量 EMBEDDING_PROVIDER 切换：local / openai

    [#11] 当 Embedding 不可用时，标记 _embedding_available = False，
    后续操作将自动降级为关键词索引模式
    """
    global _embeddings_instance, _embedding_available

    # 如果已知 Embedding 不可用，直接返回 None
    if _embedding_available is False:
        return None

    if _embeddings_instance is None:
        provider = os.environ.get("EMBEDDING_PROVIDER", EMBEDDING_PROVIDER)

        if provider == "local":
            # ===== 本地 Embedding（免费，离线，中文顶级）=====
            try:
                from langchain_huggingface import HuggingFaceEmbeddings
                model_name = os.environ.get("LOCAL_EMBEDDING_MODEL", LOCAL_EMBEDDING_MODEL)
                logger.info(f"正在加载本地 Embedding 模型: {model_name}（首次运行会自动下载，约1.2GB）...")
                _embeddings_instance = HuggingFaceEmbeddings(
                    model_name=model_name,
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"normalize_embeddings": True},
                )
                _embedding_available = True
                logger.info(f"本地 Embedding 模型加载完成: {model_name}")
            except ImportError:
                logger.error("langchain-huggingface 未安装，请运行: pip install langchain-huggingface sentence-transformers")
                logger.info("回退到 OpenAI Embedding...")
                provider = "openai"
            except Exception as e:
                logger.error(f"本地 Embedding 加载失败: {e}")
                logger.info("回退到 OpenAI Embedding...")
                provider = "openai"

        if provider == "openai":
            # ===== 云端 API Embedding（消耗额度，需联网）=====
            try:
                from langchain_openai import OpenAIEmbeddings
                embedding_model = getattr(settings, 'EMBEDDING_MODEL', 'embedding-3')
                _embeddings_instance = OpenAIEmbeddings(
                    api_key=settings.LLM_API_KEY,
                    base_url=settings.LLM_BASE_URL,
                    model=embedding_model,
                )
                logger.info(f"Embedding 模型已初始化（云端API）: {embedding_model}")
                # 注意：初始化成功不代表调用成功，实际可用性在 index_document 时验证
            except Exception as e:
                logger.error(f"云端 Embedding 初始化失败: {e}")
                _embeddings_instance = None
                _embedding_available = False
                logger.warning("Embedding 不可用，系统将使用关键词索引模式（无需向量模型即可工作）")

    return _embeddings_instance


def _get_collection_name(agent_id: str = None) -> str:
    """根据 agent_id 获取 ChromaDB collection 名称

    - agent_id 为 None 或空 → 全局知识库（普通Agent模式）
    - agent_id 有值 → 智能体专属知识库
    """
    if agent_id:
        # 用 agent_id 做 collection 名，确保合法
        safe_id = agent_id.replace('-', '_').replace(' ', '_')
        return f"agent_{safe_id}"
    return GLOBAL_COLLECTION_NAME


def get_vector_store(agent_id: str = None):
    """获取 ChromaDB 向量数据库实例（按 agent_id 隔离）

    Args:
        agent_id: 智能体ID，为 None 时使用全局知识库

    每个智能体有独立的 ChromaDB collection，互不干扰。
    普通 Agent 模式使用默认的全局 collection。

    [#11] 当 Embedding 不可用时返回 None
    """
    global _embedding_available

    # 如果已知 Embedding 不可用，直接返回 None
    if _embedding_available is False:
        return None

    cache_key = agent_id or "__global__"

    if cache_key not in _vector_store_cache:
        embeddings = get_embeddings()
        if embeddings is None:
            return None

        collection_name = _get_collection_name(agent_id)
        try:
            vs = Chroma(
                collection_name=collection_name,
                persist_directory=settings.CHROMA_DIR,
                embedding_function=embeddings,
            )
            _vector_store_cache[cache_key] = vs
            logger.info(f"ChromaDB 已连接: collection={collection_name}, agent_id={agent_id}")
        except Exception as e:
            logger.error(f"ChromaDB 连接失败: {e}")
            if _is_embedding_error(e):
                _embedding_available = False
                logger.warning("Embedding 不可用，系统将使用关键词索引模式")
            return None

    return _vector_store_cache.get(cache_key)


def reset_vector_store():
    """重置向量数据库单例（配置变更时调用）"""
    global _embeddings_instance, _embedding_available
    _vector_store_cache.clear()
    _embeddings_instance = None
    _embedding_available = None  # 重置后重新检测
    logger.info("向量数据库单例已重置，将重新检测 Embedding 可用性")


def reindex_all_documents(agent_id: str = None):
    """重建指定知识库的所有文档索引（切换embedding模型后调用）

    当从 OpenAI Embedding 切换到本地 Embedding 时，
    旧向量数据的维度不同，需要删除旧collection并重新索引。

    [#11] 同时支持向量模式和关键词模式

    Args:
        agent_id: 智能体ID，为None时重建全局知识库

    Returns:
        dict: 包含重建结果
    """
    import chromadb
    collection_name = _get_collection_name(agent_id)

    try:
        # 1. 收集所有文档来源（ChromaDB + 关键词索引 + 磁盘文件）
        document_files = set()

        # 从 ChromaDB 获取
        try:
            client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
            existing_collections = [c.name for c in client.list_collections()]
            if collection_name in existing_collections:
                collection = client.get_collection(collection_name)
                all_docs = collection.get(include=["metadatas"])
                for meta in (all_docs.get("metadatas") or []):
                    if meta and "source_file" in meta:
                        document_files.add(meta["source_file"])
                # 删除旧collection
                client.delete_collection(collection_name)
                logger.info(f"已删除旧collection: {collection_name}")
        except Exception as e:
            logger.warning(f"从ChromaDB获取文档列表失败: {e}")

        # 从关键词索引获取
        keyword_docs = _load_keyword_index(agent_id)
        for entry in keyword_docs:
            if entry.get("source_file"):
                document_files.add(entry["source_file"])

        # 从磁盘扫描
        if agent_id:
            scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
        else:
            scan_dir = settings.DOCUMENTS_DIR
        if os.path.exists(scan_dir):
            for fname in os.listdir(scan_dir):
                ext = os.path.splitext(fname)[1].lower()
                if ext in {'.pdf', '.txt', '.docx'}:
                    file_path = os.path.join(scan_dir, fname)
                    if os.path.isfile(file_path):
                        document_files.add(fname)

        # 2. 清除缓存
        cache_key = agent_id or "__global__"
        if cache_key in _vector_store_cache:
            del _vector_store_cache[cache_key]

        # 3. 清除旧关键词索引
        keyword_index_path = _get_keyword_index_path(agent_id)
        if os.path.exists(keyword_index_path):
            os.remove(keyword_index_path)

        # 4. 重新索引所有文档
        reindexed = []
        failed = []
        for filename in document_files:
            # 查找文件路径（可能在agent子目录中）
            if agent_id:
                file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
            else:
                file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

            if not os.path.exists(file_path):
                # 尝试全局目录
                file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

            if os.path.exists(file_path):
                try:
                    result = index_document(file_path, filename, agent_id=agent_id)
                    reindexed.append(filename)
                    logger.info(f"重新索引成功: {filename}, {result.get('chunks', 0)} 个分块")
                except Exception as e:
                    failed.append(f"{filename}: {str(e)}")
                    logger.error(f"重新索引失败: {filename}, {e}")
            else:
                failed.append(f"{filename}: 文件不存在")

        mode_str = "向量" if _embedding_available else "关键词"
        return {
            "status": "success",
            "collection": collection_name,
            "indexing_mode": mode_str,
            "documents_found": len(document_files),
            "reindexed": len(reindexed),
            "failed": failed,
            "message": f"知识库重建完成（{mode_str}模式）: 找到{len(document_files)}个文档，成功索引{len(reindexed)}个" + (f"，失败{len(failed)}个" if failed else "")
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"重建知识库失败: {str(e)}"
        }


def load_document(file_path: str) -> list:
    """
    根据文件类型加载文档
    支持：PDF、TXT、DOCX
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
    elif ext == ".docx":
        loader = Docx2txtLoader(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 PDF/TXT/DOCX")

    return loader.load()


def split_documents(docs: list, chunk_size: int = 500, chunk_overlap: int = 100) -> list:
    """
    文档分块
    - chunk_size: 每块最大字符数
    - chunk_overlap: 块间重叠字符数（保证上下文连续性）
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    # [#10] 给每个 chunk 分配唯一 ID，用于引用溯源
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
    return chunks


def index_document(file_path: str, filename: str = None, agent_id: str = None) -> dict:
    """
    完整的文档索引流程：加载 → 分块 → 索引存储

    [#11] 智能降级策略：
    - 当 Embedding 可用时：使用 ChromaDB 向量索引（语义搜索 + 关键词搜索）
    - 当 Embedding 不可用时：自动降级为关键词索引（纯关键词搜索）
    - 降级过程完全自动，对上层调用者透明

    Returns:
        dict: 包含分块数量和状态信息
    """
    global _embedding_available

    if filename is None:
        filename = os.path.basename(file_path)

    logger.info(f"开始索引文档: {filename}")

    # 1. 加载文档
    docs = load_document(file_path)

    # 2. 给文档添加元数据
    for doc in docs:
        doc.metadata["source_file"] = filename

    # 3. 分块
    chunks = split_documents(docs)

    if not chunks:
        return {
            "filename": filename,
            "chunks": 0,
            "status": "success",
            "message": f"文档 {filename} 内容为空，无需索引",
        }

    total_chunks = len(chunks)

    # ===== [#11] 智能索引：向量优先，关键词降级 =====

    # 情况1：已知 Embedding 不可用 → 直接使用关键词索引
    if _embedding_available is False:
        _add_chunks_to_keyword_index(chunks, filename, agent_id)
        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "keyword",
            "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding API 不可用，已自动切换为关键词搜索",
        }

    # 情况2：尝试向量索引
    try:
        vector_store = get_vector_store(agent_id=agent_id)

        if vector_store is None:
            # Embedding 不可用，降级为关键词索引
            _add_chunks_to_keyword_index(chunks, filename, agent_id)
            return {
                "filename": filename,
                "chunks": total_chunks,
                "status": "success",
                "indexing_mode": "keyword",
                "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding 不可用，已自动切换为关键词搜索",
            }

        # 尝试分批向量化并存储
        batch_count = (total_chunks + EMBEDDING_BATCH_SIZE - 1) // EMBEDDING_BATCH_SIZE

        for i in range(batch_count):
            start = i * EMBEDDING_BATCH_SIZE
            end = min(start + EMBEDDING_BATCH_SIZE, total_chunks)
            batch = chunks[start:end]
            try:
                vector_store.add_documents(batch)
            except Exception as e:
                error_str = str(e)

                if _is_embedding_error(e):
                    # [#11] Embedding 不可用 → 自动降级为关键词索引
                    logger.warning(f"Embedding 不可用（{error_str}），自动切换为关键词索引模式")

                    # 标记为不可用
                    _embedding_available = False

                    # 尝试回滚已写入该文档的数据
                    try:
                        collection = vector_store._collection
                        existing = collection.get(
                            where={"source_file": filename},
                            include=["metadatas"],
                        )
                        if existing.get("ids"):
                            collection.delete(ids=existing["ids"])
                    except Exception:
                        pass

                    # 降级：将全部分块（不仅是当前批次）写入关键词索引
                    _add_chunks_to_keyword_index(chunks, filename, agent_id)

                    return {
                        "filename": filename,
                        "chunks": total_chunks,
                        "status": "success",
                        "indexing_mode": "keyword",
                        "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding API 不可用，已自动切换为关键词搜索",
                    }
                else:
                    # 非 Embedding 错误（如 ChromaDB 内部错误），仍然尝试降级
                    logger.error(f"向量化失败（非Embedding错误）: {error_str}")

                    # 回滚已写入的数据
                    try:
                        collection = vector_store._collection
                        existing = collection.get(
                            where={"source_file": filename},
                            include=["metadatas"],
                        )
                        if existing.get("ids"):
                            collection.delete(ids=existing["ids"])
                    except Exception:
                        pass

                    # 降级为关键词索引
                    _embedding_available = False
                    _add_chunks_to_keyword_index(chunks, filename, agent_id)

                    return {
                        "filename": filename,
                        "chunks": total_chunks,
                        "status": "success",
                        "indexing_mode": "keyword",
                        "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- 向量化失败，已自动切换为关键词搜索",
                    }

        # 向量索引成功
        logger.info(f"文档索引完成（向量模式）: {filename}, 共 {total_chunks} 个分块")
        _embedding_available = True

        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "vector",
            "message": f"文档 {filename} 已成功索引（向量模式，共 {total_chunks} 个分块，分 {batch_count} 批写入）",
        }

    except Exception as e:
        # 整个向量流程异常，降级为关键词索引
        logger.error(f"向量索引流程异常: {e}")
        _embedding_available = False
        _add_chunks_to_keyword_index(chunks, filename, agent_id)

        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "keyword",
            "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- 向量索引异常，已自动切换为关键词搜索",
        }


# ===== [#9] 混合检索 + 重排序 =====

def _bm25_keyword_search(query: str, top_k: int = 10, agent_id: str = None) -> list[dict]:
    """
    [#9] BM25 风格的关键词检索（简化版）
    从 ChromaDB 获取全量文档，按关键词匹配度排序
    作为向量检索的补充，提升关键词精确匹配场景的召回率

    [#11] 当 ChromaDB 不可用时返回空列表（由 _search_keyword_index 替代）
    """
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is None:
        return []

    try:
        collection = vector_store._collection
        # 获取所有文档
        all_docs = collection.get(include=["documents", "metadatas"])

        if not all_docs.get("ids"):
            return []

        query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
        query_terms = query_terms - _STOPWORDS

        scored = []
        for i, doc_id in enumerate(all_docs["ids"]):
            content = all_docs["documents"][i] or ""
            metadata = all_docs["metadatas"][i] or {}

            # 计算关键词匹配分
            content_lower = content.lower()
            match_count = sum(1 for term in query_terms if term in content_lower)
            if match_count == 0:
                continue

            # TF 近似：关键词出现次数 / 文档长度
            tf_score = match_count / max(len(content), 1) * 1000

            scored.append({
                "content": content,
                "source": metadata.get("source_file", "未知来源"),
                "chunk_index": metadata.get("chunk_index", -1),
                "bm25_score": tf_score,
                "id": doc_id,
            })

        # 按 BM25 分排序
        scored.sort(key=lambda x: x["bm25_score"], reverse=True)
        return scored[:top_k]

    except Exception as e:
        logger.warning(f"BM25 关键词检索失败: {e}")
        return []


def _reciprocal_rank_fusion(vector_results: list[dict], keyword_results: list[dict], k: int = 60) -> list[dict]:
    """
    [#9] 倒数排名融合（Reciprocal Rank Fusion）
    将向量检索和关键词检索的结果融合，按融合分数排序

    RRF公式: score = 1/(k + rank_vector) + 1/(k + rank_keyword)
    """
    fused_scores = {}

    # 向量检索结果
    for rank, item in enumerate(vector_results):
        content_key = item["content"][:200]  # 用内容前200字作为唯一标识
        if content_key not in fused_scores:
            fused_scores[content_key] = {**item, "rrf_score": 0}
        fused_scores[content_key]["rrf_score"] += 1.0 / (k + rank + 1)
        # 保留向量相似度
        if "relevance_score" not in fused_scores[content_key]:
            fused_scores[content_key]["relevance_score"] = item.get("relevance_score", 0)

    # 关键词检索结果
    for rank, item in enumerate(keyword_results):
        content_key = item["content"][:200]
        if content_key not in fused_scores:
            fused_scores[content_key] = {
                "content": item["content"],
                "source": item.get("source", "未知来源"),
                "chunk_index": item.get("chunk_index", -1),
                "relevance_score": 0,
                "rrf_score": 0,
            }
        fused_scores[content_key]["rrf_score"] += 1.0 / (k + rank + 1)
        # 如果有 BM25 分，补充
        if "bm25_score" in item and "bm25_score" not in fused_scores[content_key]:
            fused_scores[content_key]["bm25_score"] = item["bm25_score"]

    # 按 RRF 分排序
    results = sorted(fused_scores.values(), key=lambda x: x["rrf_score"], reverse=True)
    return results


def search_documents(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """
    [#9] 混合检索：向量语义检索 + BM25关键词检索 + RRF融合
    [#10] 引用溯源：返回结果标注文档名 + 段落位置
    [#11] 自动降级：Embedding 不可用时仅使用关键词检索

    Args:
        query: 用户查询
        top_k: 返回最相关的 K 个结果
        agent_id: 智能体ID

    Returns:
        list[dict]: 检索结果列表
    """
    # ===== [#11] 根据索引模式选择检索策略 =====

    if _embedding_available is False:
        # 关键词模式：仅使用关键词索引检索
        logger.info(f"关键词模式检索: query='{query[:50]}...', agent_id={agent_id}")
        results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
        if not results:
            # 关键词索引无结果，尝试从磁盘文件全文搜索
            results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
        return results

    # ===== 向量模式：混合检索 =====

    # 1. 向量语义检索（按 agent_id 隔离）
    vector_store = get_vector_store(agent_id=agent_id)
    vector_results_raw = []

    if vector_store is not None:
        try:
            vector_results_raw = vector_store.similarity_search_with_score(query, k=top_k * 3)  # 多取用于融合
        except Exception as e:
            error_str = str(e)
            logger.warning(f"向量检索失败: {e}")
            if _is_embedding_error(e):
                logger.warning(f"Embedding API 不可用，自动切换为关键词检索模式")
                _embedding_available = False
                # 降级为关键词检索
                results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
                if not results:
                    results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
                return results
            vector_results_raw = []
    else:
        # vector_store 为 None，降级为关键词模式
        _embedding_available = False
        results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
        if not results:
            results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
        return results

    vector_results = []
    for doc, score in vector_results_raw:
        vector_results.append({
            "content": doc.page_content,
            "source": doc.metadata.get("source_file", "未知来源"),
            "chunk_index": doc.metadata.get("chunk_index", -1),
            "relevance_score": round(1 - score, 4),
        })

    # 2. BM25 关键词检索（按 agent_id 隔离）
    keyword_results = _bm25_keyword_search(query, top_k=top_k * 3, agent_id=agent_id)

    # 3. RRF 融合
    if keyword_results:
        fused_results = _reciprocal_rank_fusion(vector_results, keyword_results)
    else:
        # 关键词检索失败，直接用向量结果
        fused_results = vector_results

    # 4. 格式化输出（取 top_k）
    formatted = []
    for r in fused_results[:top_k]:
        formatted.append({
            "content": r["content"],
            "source": r.get("source", "未知来源"),
            "chunk_index": r.get("chunk_index", -1),
            "relevance_score": r.get("relevance_score", round(r.get("rrf_score", 0), 4)),
        })

    return formatted


def _search_disk_files(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """[#11] 磁盘文件全文搜索（关键词索引的补充）

    当关键词索引中也没有匹配结果时，直接读取磁盘上的文档文件做搜索。
    这是最后的兜底方案，确保即使没有任何索引，用户也能查到文档内容。

    Args:
        query: 查询文本
        top_k: 返回结果数
        agent_id: 智能体ID

    Returns:
        list[dict]: 搜索结果
    """
    if agent_id:
        scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
    else:
        scan_dir = settings.DOCUMENTS_DIR

    if not os.path.exists(scan_dir):
        return []

    query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
    query_terms = query_terms - _STOPWORDS

    if not query_terms:
        return []

    scored = []
    for fname in os.listdir(scan_dir):
        ext = os.path.splitext(fname)[1].lower()
        if ext not in {'.txt', '.docx', '.pdf'}:
            continue
        file_path = os.path.join(scan_dir, fname)
        if not os.path.isfile(file_path):
            continue

        try:
            docs = load_document(file_path)
            for doc in docs:
                content = doc.page_content
                content_lower = content.lower()
                match_count = sum(1 for term in query_terms if term in content_lower)
                if match_count > 0:
                    term_coverage = match_count / max(len(query_terms), 1)
                    tf_score = match_count / max(len(content), 1) * 1000
                    combined_score = tf_score * 0.6 + term_coverage * 100 * 0.4
                    scored.append({
                        "content": content[:2000],  # 限制长度避免过大
                        "source": fname,
                        "chunk_index": 0,
                        "relevance_score": round(combined_score, 4),
                    })
        except Exception:
            continue

    scored.sort(key=lambda x: x["relevance_score"], reverse=True)
    return scored[:top_k]


def get_document_content(filename: str, agent_id: str = None) -> dict:
    """获取知识库中指定文档的完整内容（从磁盘原始文件读取，不依赖向量搜索）

    与 search_documents 不同，此函数返回文档的完整文本内容，
    而不是分块后的片段。用于文档修改前获取完整内容。

    Args:
        filename: 文档文件名（含扩展名）
        agent_id: 智能体ID（用于验证文档归属，不参与向量搜索）

    Returns:
        dict: 包含文档完整内容、状态信息
    """
    # 查找文件路径（可能在agent子目录中）
    if agent_id:
        file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
        if not os.path.exists(file_path):
            file_path = os.path.join(settings.DOCUMENTS_DIR, filename)
    else:
        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

    if not os.path.exists(file_path):
        return {
            "filename": filename,
            "status": "not_found",
            "content": "",
            "message": f"文档 {filename} 在服务器上未找到",
        }

    try:
        docs = load_document(file_path)
        full_content = "\n".join([doc.page_content for doc in docs])

        if not full_content.strip():
            return {
                "filename": filename,
                "status": "empty",
                "content": "",
                "message": f"文档 {filename} 内容为空",
            }

        return {
            "filename": filename,
            "status": "success",
            "content": full_content,
            "char_count": len(full_content),
            "message": f"成功获取文档 {filename} 的完整内容（共 {len(full_content)} 字符）",
        }
    except Exception as e:
        return {
            "filename": filename,
            "status": "error",
            "content": "",
            "message": f"读取文档失败: {str(e)}",
        }


def list_indexed_documents(agent_id: str = None) -> list[str]:
    """列出知识库中所有已索引的文档（按 agent_id 隔离）

    [#11] 同时检查向量索引和关键词索引，合并结果
    """
    sources = set()

    # 1. 从 ChromaDB 获取（如果可用）
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            all_docs = collection.get(include=["metadatas"])
            for meta in all_docs["metadatas"]:
                if meta and "source_file" in meta:
                    sources.add(meta["source_file"])
        except Exception:
            pass

    # 2. 从关键词索引获取
    keyword_docs = _load_keyword_index(agent_id)
    for entry in keyword_docs:
        if entry.get("source_file"):
            sources.add(entry["source_file"])

    return sorted(list(sources))


def update_document(filename: str, new_content: str, agent_id: str = None, async_reindex: bool = False) -> dict:
    """
    修改知识库中已有文档的内容
    流程：删除旧的向量分块 → 用新内容覆盖原文件 → 重新索引

    [#11] 同时更新向量索引和关键词索引

    Args:
        filename: 要修改的文档文件名
        new_content: 新的文档内容（纯文本）
        agent_id: 智能体ID
        async_reindex: 是否异步重索引

    Returns:
        dict: 包含修改状态和详细信息
    """
    # 1. 检查文件是否存在（可能在agent子目录中）
    if agent_id:
        file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
        if not os.path.exists(file_path):
            file_path = os.path.join(settings.DOCUMENTS_DIR, filename)
    else:
        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

    if not os.path.exists(file_path):
        return {
            "filename": filename,
            "status": "not_found",
            "message": f"文档 {filename} 在服务器上未找到",
        }

    # 2. 删除旧的索引数据
    chunks_deleted = 0

    # 从 ChromaDB 删除
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            results = collection.get(
                where={"source_file": filename},
                include=["metadatas"],
            )
            chunk_ids = results.get("ids", [])
            if chunk_ids:
                collection.delete(ids=chunk_ids)
                chunks_deleted = len(chunk_ids)
        except Exception as e:
            logger.warning(f"删除旧向量分块时出错: {e}")

    # 从关键词索引删除
    keyword_deleted = _delete_from_keyword_index(filename, agent_id)

    # 3. 用新内容覆盖原文件
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".txt":
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(new_content)
        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                for line in new_content.split("\n"):
                    doc.add_paragraph(line)
                doc.save(file_path)
            except ImportError:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(new_content)
        elif ext == ".pdf":
            txt_path = file_path.rsplit('.', 1)[0] + '.txt'
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(new_content)
            os.remove(file_path)
            filename = filename.rsplit('.', 1)[0] + '.txt'
            file_path = txt_path
        else:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(new_content)
    except Exception as e:
        return {
            "filename": filename,
            "status": "error",
            "message": f"写入文件失败: {str(e)}",
        }

    # 4. 重新索引
    if async_reindex:
        import threading
        def _background_reindex(fp, fn, aid):
            try:
                index_result = index_document(fp, fn, agent_id=aid)
                logger.info(f"后台重索引完成: {fn}, {index_result.get('chunks', 0)} 个分块, 模式={index_result.get('indexing_mode', 'unknown')}")
            except Exception as e:
                logger.error(f"后台重索引失败: {fn}, {e}")

        thread = threading.Thread(target=_background_reindex, args=(file_path, filename, agent_id), daemon=True)
        thread.start()

        return {
            "filename": filename,
            "status": "success",
            "chunks_deleted": chunks_deleted,
            "keyword_entries_deleted": keyword_deleted,
            "chunks_indexed": "后台索引中",
            "message": f"文档 {filename} 已成功修改（删除 {chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，新内容正在后台索引中）",
        }
    else:
        try:
            index_result = index_document(file_path, filename, agent_id=agent_id)
        except Exception as e:
            return {
                "filename": filename,
                "status": "error",
                "message": f"重新索引失败: {str(e)}",
                "chunks_deleted": chunks_deleted,
            }

        mode = index_result.get("indexing_mode", "unknown")
        return {
            "filename": filename,
            "status": "success",
            "chunks_deleted": chunks_deleted,
            "keyword_entries_deleted": keyword_deleted,
            "chunks_indexed": index_result.get("chunks", 0),
            "indexing_mode": mode,
            "message": f"文档 {filename} 已成功修改（{mode}模式，删除 {chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，重新索引 {index_result.get('chunks', 0)} 个新分块）",
        }


def delete_document(filename: str, agent_id: str = None) -> dict:
    """
    从知识库中删除指定文档
    包括：从 ChromaDB 删除向量分块 + 从关键词索引删除 + 删除原始文件

    [#11] 同时清理向量索引和关键词索引

    Args:
        filename: 要删除的文档文件名
        agent_id: 智能体ID

    Returns:
        dict: 包含删除状态和详细信息
    """
    chunks_deleted = 0
    found_in_any = False

    # 1. 从 ChromaDB 删除
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            results = collection.get(
                where={"source_file": filename},
                include=["metadatas"],
            )
            chunk_ids = results.get("ids", [])
            if chunk_ids:
                found_in_any = True
                collection.delete(ids=chunk_ids)
                chunks_deleted = len(chunk_ids)
        except Exception as e:
            logger.warning(f"从 ChromaDB 删除失败: {e}")

    # 2. 从关键词索引删除
    keyword_deleted = _delete_from_keyword_index(filename, agent_id)
    if keyword_deleted > 0:
        found_in_any = True

    # 3. 删除原始文件（查找可能的位置）
    file_deleted = False
    possible_paths = []

    if agent_id:
        possible_paths.append(os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename))
    possible_paths.append(os.path.join(settings.DOCUMENTS_DIR, filename))

    for file_path in possible_paths:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                file_deleted = True
                found_in_any = True
                break
            except Exception as e:
                return {
                    "filename": filename,
                    "chunks_deleted": chunks_deleted,
                    "keyword_entries_deleted": keyword_deleted,
                    "file_deleted": False,
                    "status": "partial",
                    "message": f"索引已删除，但原始文件删除失败: {str(e)}",
                }

    if not found_in_any:
        return {
            "filename": filename,
            "status": "not_found",
            "message": f"文档 {filename} 在知识库中未找到",
        }

    return {
        "filename": filename,
        "chunks_deleted": chunks_deleted,
        "keyword_entries_deleted": keyword_deleted,
        "file_deleted": file_deleted,
        "status": "success",
        "message": f"文档 {filename} 已成功删除（{chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，原始文件{'已删除' if file_deleted else '不存在'}）",
    }


def delete_agent_collection(agent_id: str) -> dict:
    """删除智能体的整个知识库 collection

    删除智能体时调用，清理 ChromaDB 中该智能体专属的 collection。
    同时清理关键词索引文件。

    Args:
        agent_id: 智能体ID

    Returns:
        dict: 包含删除状态和详细信息
    """
    if not agent_id:
        return {"status": "error", "message": "agent_id 不能为空"}

    import chromadb
    collection_name = _get_collection_name(agent_id)

    try:
        # 1. 删除 ChromaDB collection
        client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
        existing_collections = [c.name for c in client.list_collections()]
        if collection_name in existing_collections:
            client.delete_collection(collection_name)
            logger.info(f"已删除智能体 ChromaDB collection: {collection_name}")

        # 2. 清理缓存
        cache_key = agent_id or "__global__"
        if cache_key in _vector_store_cache:
            del _vector_store_cache[cache_key]

        # 3. 删除关键词索引文件
        keyword_index_path = _get_keyword_index_path(agent_id)
        keyword_deleted = False
        if os.path.exists(keyword_index_path):
            try:
                os.remove(keyword_index_path)
                keyword_deleted = True
                logger.info(f"已删除智能体关键词索引: {keyword_index_path}")
            except Exception as e:
                logger.warning(f"删除关键词索引文件失败: {e}")

        print(f"[DEBUG-知识库] 已删除智能体 collection: {collection_name}, 关键词索引: {'已删除' if keyword_deleted else '无'}")
        return {"status": "success", "message": f"智能体知识库 {collection_name} 已删除（ChromaDB + 关键词索引）"}
    except Exception as e:
        print(f"[DEBUG-知识库] 删除智能体 collection 失败: {e}")
        return {"status": "error", "message": f"删除知识库失败: {str(e)}"}


def export_document_as_docx(content: str, filename: str, title: str = "") -> dict:
    """
    将文本内容导出为 .docx 文件，保存到下载目录，供用户下载

    Args:
        content: 文档内容（纯文本）
        filename: 输出文件名（含扩展名）
        title: 文档标题（可选，将作为文档第一行加粗显示）

    Returns:
        dict: 包含导出状态和文件路径
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # 回退方案：保存为 .txt
        txt_filename = filename.rsplit('.', 1)[0] + '.txt'
        file_path = os.path.join(settings.DOCUMENTS_DIR, txt_filename)
        with open(file_path, "w", encoding="utf-8") as f:
            if title:
                f.write(f"{title}\n{'=' * len(title.encode('gbk', errors='replace'))}\n\n")
            f.write(content)
        return {
            "status": "success",
            "filename": txt_filename,
            "file_path": file_path,
            "message": f"文档已导出为 {txt_filename}（python-docx 未安装，回退为 txt 格式）",
        }

    try:
        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)
        doc = DocxDocument()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(11)

        # 添加标题
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc_title = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ')
            heading = doc.add_heading(doc_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 按段落写入内容
        lines = content.split('\n')
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                doc.add_paragraph('')
                continue

            if line_stripped.startswith('# ') and not line_stripped.startswith('## '):
                text = line_stripped[2:].strip()
                p = doc.add_heading(text, level=2)
            elif line_stripped.startswith('## '):
                text = line_stripped[3:].strip()
                p = doc.add_heading(text, level=2)
            elif line_stripped.startswith('### '):
                text = line_stripped[4:].strip()
                p = doc.add_heading(text, level=3)
            elif line_stripped.startswith('#### '):
                text = line_stripped[5:].strip()
                p = doc.add_heading(text, level=4)
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:].strip()
                text = _clean_markdown_formatting(text)
                doc.add_paragraph(text, style='List Bullet')
            elif line_stripped.startswith('|') and '|' in line_stripped[1:]:
                p = doc.add_paragraph(line_stripped)
                p.style = doc.styles['Normal']
            else:
                text = _clean_markdown_formatting(line_stripped)
                doc.add_paragraph(text)

        doc.save(file_path)

        return {
            "status": "success",
            "filename": filename,
            "file_path": file_path,
            "message": f"文档已导出为 {filename}",
        }
    except Exception as e:
        txt_filename = filename.rsplit('.', 1)[0] + '.txt'
        file_path = os.path.join(settings.DOCUMENTS_DIR, txt_filename)
        with open(file_path, "w", encoding="utf-8") as f:
            if title:
                f.write(f"{title}\n\n")
            f.write(content)
        return {
            "status": "success",
            "filename": txt_filename,
            "file_path": file_path,
            "message": f"文档已导出为 {txt_filename}（docx 生成失败: {str(e)}，回退为 txt 格式）",
        }


def _clean_markdown_formatting(text: str) -> str:
    """清理Markdown格式标记，转为纯文本"""
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def list_all_collections() -> list[dict]:
    """列出 ChromaDB 中所有的 collection 及其文档数（诊断用）"""
    import chromadb
    try:
        client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
        collections = client.list_collections()
        result = []
        for c in collections:
            try:
                count = c.count()
            except:
                count = -1
            result.append({"name": c.name, "count": count})

        # [#11] 同时列出关键词索引信息
        keyword_info = []
        if os.path.exists(KEYWORD_INDEX_DIR):
            for fname in os.listdir(KEYWORD_INDEX_DIR):
                if fname.startswith("index_") and fname.endswith(".json"):
                    fpath = os.path.join(KEYWORD_INDEX_DIR, fname)
                    try:
                        with open(fpath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        keyword_info.append({"name": fname, "count": len(data), "type": "keyword"})
                    except Exception:
                        pass

        result.extend(keyword_info)
        return result
    except Exception as e:
        return [{"name": "error", "count": 0, "message": str(e)}]
